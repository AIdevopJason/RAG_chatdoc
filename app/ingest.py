from pathlib import Path
from pypdf import PdfReader
from docx import Document


def read_txt(file_path: Path) -> list[dict]:
    text = file_path.read_text(encoding="utf-8")
    return [
        {
            "source": file_path.name,
            "page": 1,
            "text": text.strip(),
            "char_count": len(text.strip()),
            "file_type": "txt",
        }
    ]


def read_pdf(file_path: Path) -> list[dict]:
    reader = PdfReader(str(file_path))
    pages = []

    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()

        pages.append(
            {
                "source": file_path.name,
                "page": i,
                "text": text,
                "char_count": len(text),
                "file_type": "pdf",
            }
        )

    return pages


def read_docx(file_path: Path) -> list[dict]:
    document = Document(str(file_path))
    text_parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_parts.append(" | ".join(cells))

    text = "\n\n".join(text_parts).strip()
    return [
        {
            "source": file_path.name,
            "page": 1,
            "text": text,
            "char_count": len(text),
            "file_type": "docx",
        }
    ]


def load_document(file_path: str) -> list[dict]:
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".txt":
        return read_txt(path)
    elif suffix == ".pdf":
        return read_pdf(path)
    elif suffix == ".docx":
        return read_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


if __name__ == "__main__":
    sample_file = "../data/raw/AG网络迁移计划.docx"  # 改成你的文件名

    try:
        docs = load_document(sample_file)

        print(f"Loaded {len(docs)} page(s)/document chunk(s).\n")

        for doc in docs[:3]:
            print("=" * 80)
            print(f"Source: {doc['source']}")
            print(f"Page: {doc['page']}")
            print(f"Characters: {doc['char_count']}")
            print("Preview:")
            print(doc["text"][:500] if doc["text"] else "[EMPTY PAGE]")
            print()

    except Exception as e:
        print(f"Error: {e}")